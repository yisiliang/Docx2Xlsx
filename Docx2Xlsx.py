import docx
import re

import openpyxl

from openpyxl.worksheet.copier import WorksheetCopy


def sFullToHalf(s):
    str = ''
    for char in s:
        num = ord(char)
        if num == 0x3000:
            num = 32
        elif 0xFF01 <= num <= 0xFF5E:
            num -= 0xfee0
        str = str + chr(num)
    return str


def get_column_type(s):
    s = sFullToHalf(s)
    s = sTrimSpace(s)
    pos = s.find('(')
    s = s[:pos]
    return s.upper()


def get_column_len(s):
    s = sFullToHalf(s)
    s = sTrimSpace(s)
    pos1 = s.find('(')
    pos2 = s.find(')')
    s = s[pos1 + 1:pos2]
    return s


def sTrimSpace(s):
    str = ''
    for char in s:
        if char != ' ':
            str = str + char
    return str


document = docx.Document('from.docx')

textlist = [paragraph.text for paragraph in document.paragraphs]
tabList = document.tables

nameDict = dict()
detailDict = dict()

count = 0
for string in textlist:
    string = sTrimSpace(sFullToHalf(string))
    if re.match('.+表\([a-zA-Z0-9_]+\)\Z', string) is not None:
        start_pos = string.find('表(')
        start_pos = start_pos + 2
        tabName = string[start_pos:-1]
        tabName = tabName.upper()
        tabCHNName = string[:start_pos - 1]

        print('tabName = ' + tabName + ', tabCHNName = ' + tabCHNName)
        nameDict[tabName] = tabCHNName
        count = count + 1

print(count)

workbook = openpyxl.load_workbook(
    'temple.xlsx')

menuSheet = workbook.get_sheet_by_name('目录')
sysSheet = workbook.get_sheet_by_name('SAMPLE')

menuRow = 3

tablePos = 0

for key in nameDict:

    sheetNames = workbook.get_sheet_names()
    sheetIndex = len(sheetNames)
    newSheet = workbook.create_sheet(key, sheetIndex + 1)
    copy = openpyxl.worksheet.copier.WorksheetCopy(sysSheet, newSheet)
    WorksheetCopy.copy_worksheet(copy)

    copySheet = workbook.get_sheet_by_name(key)
    copySheet.cell(row=2, column=3).value = key
    copySheet.cell(row=2, column=5).value = nameDict[key]

    menuSheet.cell(row=menuRow, column=4).value = '= HYPERLINK("#' + key + '!A1","' + key + '")'
    menuSheet.cell(row=menuRow, column=5).value = nameDict[key]
    menuRow = menuRow + 1

    table = tabList[tablePos]
    tablePos = tablePos + 1

    print(str(tablePos) + '/' + str(len(tabList)) + ': tabName = ' + key + ', tabCHNName = ' + nameDict[key])

    rowCnt = 0
    rowStart = 3
    for row in table.rows:
        rowCnt = rowCnt + 1
        if rowCnt > 1:
            if len(row.cells) == 6:
                # 字段名
                copySheet.cell(row=rowStart, column=3).value = row.cells[1].text.rstrip().upper()
                # 中文名
                copySheet.cell(row=rowStart, column=4).value = row.cells[2].text.rstrip().upper()
                # 类型
                copySheet.cell(row=rowStart, column=5).value = get_column_type(row.cells[3].text.rstrip())
                # 长度
                lenstr = get_column_len(row.cells[3].text.rstrip())
                if lenstr is not None:
                    try:
                        lenint = int(lenstr.rstrip())
                        copySheet.cell(row=rowStart, column=6).value = lenint
                    except ValueError:
                        print(lenstr + ' is not number')

                # 键值
                copySheet.cell(row=rowStart, column=7).value = row.cells[4].text.rstrip().upper()
                # 空值
                copySheet.cell(row=rowStart, column=8).value = ''
                # 字段说明
                copySheet.cell(row=rowStart, column=9).value = row.cells[5].text.rstrip()

            if len(row.cells) == 8:
                # 字段名
                copySheet.cell(row=rowStart, column=3).value = row.cells[1].text.rstrip().upper()
                # 中文名
                copySheet.cell(row=rowStart, column=4).value = row.cells[2].text.rstrip().upper()
                # 类型
                copySheet.cell(row=rowStart, column=5).value = row.cells[3].text.rstrip().upper()
                # 长度
                lenstr = row.cells[4].text.rstrip()
                if lenstr is not None:
                    try:
                        lenint = int(lenstr.rstrip())
                        copySheet.cell(row=rowStart, column=6).value = lenint
                    except ValueError:
                        print(lenstr + ' is not number')
                # 键值
                copySheet.cell(row=rowStart, column=7).value = row.cells[5].text.rstrip().upper()
                # 空值
                copySheet.cell(row=rowStart, column=8).value = row.cells[6].text.rstrip().upper()
                # 字段说明
                copySheet.cell(row=rowStart, column=9).value = row.cells[7].text.rstrip()
        rowStart = rowStart + 1

workbook.save(
    'new.xlsx')
